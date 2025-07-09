import os
import json
import time
import base64
import binascii
import secrets
import logging
import inspect
import traceback
from functools import wraps
from datetime import datetime, timedelta
from urllib.parse import quote
from django.core.mail import send_mail
from django.http import HttpResponse
from django.template.loader import get_template
from xhtml2pdf import pisa
from django.http import FileResponse
from docx import Document
from .forms import SettlementCalculatorForm, PaymentDirectionForm, PaymentDirectionLineItemForm, RatesAdjustmentForm
import io
from decimal import Decimal, InvalidOperation
from .models import RatesAdjustment
from decimal import Decimal, ROUND_HALF_UP


import pytz
from django import forms
from django.conf import settings
from django.urls import reverse, reverse_lazy
from django.http import JsonResponse, HttpResponse, HttpResponseRedirect
from django.shortcuts import render, redirect, get_object_or_404
from django.views.decorators.csrf import csrf_protect, csrf_exempt
from django.views.decorators.http import require_POST
from django.utils.decorators import method_decorator
from django.utils.timezone import now, localtime, is_naive, get_current_timezone, make_aware
from django.contrib import messages
from django.contrib.auth import login, authenticate, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib.auth.views import LoginView as DjangoLoginView, PasswordResetView
from django.db import transaction
from django.db.models import Q

from django_otp import login as otp_login
from django_otp.decorators import otp_required
from django_otp.plugins.otp_totp.models import TOTPDevice

from two_factor.utils import default_device
from two_factor.forms import AuthenticationTokenForm, BackupTokenForm, DeviceValidationForm, TOTPDeviceForm
from two_factor.views import LoginView as TwoFactorLoginView
from two_factor.views.core import SetupView

from .models import (
    Instruction,
    Solicitor,
    Document as DocumentModel,
    Firm,
    ChatMessage,
    PaymentDirection,
    PaymentDirectionLineItem
)
from .forms import (
    LoginForm,
    WelcomeStepForm,
    ValidationStepForm,
    InstructionForm,
    DocumentUploadForm,
    CustomTOTPDeviceForm,
)

# Logger setup
logger = logging.getLogger(__name__)
logger.debug("🚀 Logger initialized and views.py loaded")


class SettlexTwoFactorLoginView(TwoFactorLoginView):
    template_name = "two_factor/login.html"

    def get_form_list(self):
        return {
            'auth': LoginForm,  # ✅ Your custom LoginForm
            'token': AuthenticationTokenForm,
            'backup': BackupTokenForm,
        }

    def dispatch(self, request, *args, **kwargs):
        logger.debug("🚀 SettlexTwoFactorLoginView: dispatch triggered.")
        return super().dispatch(request, *args, **kwargs)

@method_decorator(login_required, name='dispatch')
class SettlexTwoFactorSetupView(SetupView):
    form_list = (
        ('welcome', WelcomeStepForm),
        ('generator', CustomTOTPDeviceForm),
        ('validation', ValidationStepForm),
    )

    template_name = 'two_factor/setup.html'

    def dispatch(self, request, *args, **kwargs):
        # 🚫 Bypass 2FA wizard for superusers
        if request.user.is_superuser:
            logger.debug("🚫 Superuser detected, skipping 2FA setup.")
            return redirect('admin:index')

        # ✅ Skip setup if user already has confirmed TOTP device
        existing = TOTPDevice.objects.filter(user=request.user, confirmed=True).first()
        if existing:
            logger.debug("🔁 User already has confirmed device, redirecting to dashboard.")
            return redirect('settlements_app:my_transactions')

        return super().dispatch(request, *args, **kwargs)

    def get_form_list(self):
        form_list = super().get_form_list()
        form_list['generator'] = CustomTOTPDeviceForm
        form_list['validation'] = ValidationStepForm
        return form_list

    def get_device(self):
        """
        Custom retrieval of the TOTPDevice using extra_data stored during the wizard.
        """
        extra_data = self.storage.extra_data or {}
        device_id = extra_data.get("device_id")
        if device_id:
            try:
                return TOTPDevice.objects.get(
                    id=int(device_id), user=self.request.user)
            except TOTPDevice.DoesNotExist:
                logger.warning(
                    "❌ No TOTPDevice found for ID %s for user %s",
                    device_id,
                    self.request.user)
        return None


    def get_form(self, step=None, data=None, files=None):
        step = step or self.steps.current or self.steps.first
        form_class = self.form_list[step]
        logger.debug("📋 get_form called — step=%s, form_class=%s", step, form_class.__name__)

        kwargs = self.get_form_kwargs(step)

        # ✅ FIX 3 — Prevent recreating TOTPDevice after user has confirmed setup
        if step in ('generator', 'validation'):
            confirmed_device = TOTPDevice.objects.filter(user=self.request.user, confirmed=True).first()
            if confirmed_device:
                logger.debug("✅ Confirmed TOTPDevice already exists for user %s — skipping new device setup", self.request.user)
                if 'device' in inspect.signature(form_class).parameters:
                    kwargs['device'] = confirmed_device
                if data is None:
                    data = self.request.POST
                return form_class(data=data, files=files, **kwargs)

        device = None
        if step in ('generator', 'validation'):
            extra_data = self.storage.extra_data or {}
            device_id = extra_data.get('device_id')

            logger.debug("📦 Looking for device_id: %s in extra_data", device_id)

            if device_id:
                try:
                    device = TOTPDevice.objects.get(id=int(device_id))
                    logger.debug("📦 Reusing existing TOTPDevice: %s", device)
                except (TOTPDevice.DoesNotExist, ValueError):
                    logger.warning("⚠️ Invalid or missing device_id; creating new TOTPDevice")

            if not device:
                logger.debug("🔧 No existing device found, creating new TOTP device.")
                key = self.get_key(step)
                device = TOTPDevice.objects.create(
                    user=self.request.user,
                    confirmed=False,
                    key=key,
                    digits=6,
                )
                extra_data['device_id'] = device.id
                self.storage.extra_data = extra_data
                self.request.session[self.storage.prefix] = self.storage.data
                logger.debug("🛠 Created new TOTPDevice (ID: %s) with key: %s", device.id, device.key)

            logger.debug("📦 Device at this step: %s", device)

        # ✅ Always pass device if expected by the form
        if 'device' in inspect.signature(form_class).parameters:
            kwargs['device'] = device
            logger.debug("📦 Passing device to form: %s", device)

        if data is None:
            data = self.request.POST  # Ensure binding POST data

        return form_class(data=data, files=files, **kwargs)

    def get_context_data(self, form, **kwargs):
        step = self.steps.current or self.steps.first
        logger.debug("📦 get_context_data - Step: %s", step)

        context = super().get_context_data(form=form, **kwargs)
        logger.debug("🧭 Current step: %s", step)

        if step == 'generator':
            logger.debug("📥 Generator form raw data: %s", form.data)

            form_context = {}
            try:
                if hasattr(form, 'get_context_data'):
                    form_context = form.get_context_data()
                    logger.debug("🧬 Context from generator form: %s", form_context)
                context.update(form_context)
            except Exception:
                logger.exception("⚠️ Exception while building generator context")

            logger.debug(
                "🚨 QR Code base64 length: %s",
                len(context.get('qr_code_base64') or ''))
            logger.debug("🚨 TOTP Secret: %s", context.get('totp_secret'))
            logger.debug("📸 Generator form device: %s", getattr(form, 'device', None))
            logger.debug("🧾 Form is_bound: %s", form.is_bound)

        elif step == 'validation':
            logger.debug("📥 Validation form raw data: %s", form.data)

            form_context = {}
            try:
                if hasattr(form, 'get_context_data'):
                    form_context = form.get_context_data()
                    logger.debug("🧬 Context from validation form: %s", form_context)
                context.update(form_context)
            except Exception:
                logger.exception("⚠️ Exception while building validation context")

            logger.debug("🚨 Validation step context: %s", context)

        return context

    def get(self, request, *args, **kwargs):
        step = self.steps.current or self.steps.first
        logger.debug(f"🔍 SettlexTwoFactorSetupView: GET at step '{step}'")
        return super().get(request, *args, **kwargs)

    def post(self, request, *args, **kwargs):
        step = self.steps.current or self.steps.first
        logger.debug(f"📨 SettlexTwoFactorSetupView: POST at step '{step}'")
        logger.debug("📨 POST data: %s", request.POST)

        response = super().post(request, *args, **kwargs)

        form = self.get_form(step)
        logger.debug("🧾 Form is_bound: %s", form.is_bound)
        logger.debug("🧪 Form valid: %s", form.is_valid())
        logger.debug("🧪 Form errors: %s", form.errors)
        logger.debug(
            "🧪 Form cleaned_data: %s", getattr(
                form, 'cleaned_data', {}))

        return response

    def done(self, form_list, **kwargs):
        try:
            del self.request.session[self.session_key_name]
        except KeyError:
            pass

        device = None
        for form in form_list:
            if isinstance(form, ValidationStepForm):
                device = form.save()
                break

        if device:
            # ✅ Mark as confirmed
            device.confirmed = True

            # ✅ OPTIONAL: Assign a name (helpful for admin/debugging)
            device.name = "Primary Device"

            device.save()

            # ✅ Ensure it's the default device (this is critical to avoid redirect loop)
            from two_factor.utils import default_device
            from django_otp.plugins.otp_totp.models import TOTPDevice

            TOTPDevice.objects.filter(user=self.request.user).exclude(id=device.id).update(name=None)
            device.name = "default"
            device.save()

            # ✅ Log the device details
            logger.info("✅ 2FA setup complete for user: %s — redirecting to dashboard.", self.request.user)
            logger.debug("🔎 Default device now: %s", default_device(self.request.user))

            # ✅ Log the user in with 2FA
            otp_login(self.request, device)

        return redirect('settlements_app:my_transactions')


# Home View


def home(request):
    user = request.user
    logger.debug("Home view accessed by user: %s",
                 user.username if user.is_authenticated else "Anonymous")

    if user.is_authenticated:
        if not default_device(user):
            logger.debug(
                "✅ Authenticated but no 2FA device – redirecting to setup")
            return redirect('two_factor:setup')

        logger.debug(
            "✅ Authenticated and 2FA verified – redirecting to dashboard")
        return redirect('settlements_app:my_transactions')

    logger.debug("Rendering public home page")
    return render(request, 'settlements_app/home.html', {
        'page_title': 'Welcome to SettleX'
    })


# ✅ Logout View - Clears Session
def logout_view(request):
    logout(request)
    request.session.flush()  # ✅ Clears the session completely
    messages.success(request, "You have been logged out successfully.")
    return redirect('settlements_app:home')

# ✅ Registration View with Email Notifications


def register(request):
    firm = None

    if request.method == 'POST':
        first_name = request.POST.get('first_name', '').strip()
        last_name = request.POST.get('last_name', '').strip()
        email = request.POST.get('email', '').strip()
        password = request.POST.get('password', '').strip()
        firm_name = request.POST.get('firm_name', '').strip()
        office_phone = request.POST.get('office_phone', '').strip()
        address = request.POST.get('address', '').strip()
        postcode = request.POST.get('postcode', '').strip()
        state = request.POST.get('state', '').strip()
        mobile_phone = request.POST.get('mobile_phone', '').strip()
        profession = request.POST.get('profession', '').strip()
        law_society_number = request.POST.get('law_society_number', '').strip()
        conveyancer_license_number = request.POST.get(
            'conveyancer_license_number', '').strip()  # Fixed field name

        logger.debug("🔍 Registration attempt for email: %s", email)
        logger.debug("📋 Form data: %s", request.POST)

        # Check if the user already exists
        existing_user = User.objects.filter(email=email).first()
        if existing_user:
            if existing_user.is_active:
                messages.error(
                    request, "This email is already registered and active. Please log in.")
                logger.warning(
                    "⚠️ Email already exists and is active: %s", email)
            else:
                messages.info(
                    request,
                    "Your account is pending admin approval. You will receive an email once activated.")
                logger.info("ℹ️ Account pending approval for email: %s", email)
            return render(request,
                          'settlements_app/register.html',
                          {'message': 'Register an Account'})

        # Ensure all required fields are provided
        if not (first_name and last_name and email and password and firm_name and address and postcode and state and profession):
            messages.error(request, "Please fill in all required fields.")
            logger.warning("⚠️ Missing required fields")
            return render(request,
                          'settlements_app/register.html',
                          {'message': 'Register an Account'})

        try:
            # Check if the firm exists, if not, create it
            firm = Firm.objects.filter(name__iexact=firm_name).first()
            if not firm:
                firm = Firm.objects.create(
                    name=firm_name,
                    contact_email=email,
                    contact_number=office_phone,
                    address=address,
                    postcode=postcode,
                    state=state,
                )
                logger.info("🏢 New firm created: %s", firm.name)

            # Create the user and set to inactive (pending admin approval)
            user = User.objects.create_user(
                username=email, email=email, password=password)
            user.first_name = first_name
            user.last_name = last_name
            user.is_active = False  # Pending admin approval
            user.save()
            logger.debug("👤 User created: %s", user.username)

            # Create Solicitor profile with the correct license number
            solicitor_data = {
                'user': user,
                'instructing_solicitor': f"{first_name} {last_name}",
                'firm': firm,
                'office_phone': office_phone,
                'mobile_phone': mobile_phone,
                'profession': profession,
            }
            if profession == "solicitor":
                solicitor_data['law_society_number'] = law_society_number
                # Ensure field name matches model
                solicitor_data['conveyancer_license'] = ""
            else:
                solicitor_data['law_society_number'] = ""
                solicitor_data['conveyancer_license'] = conveyancer_license_number

            solicitor = Solicitor.objects.create(**solicitor_data)
            logger.debug(
                "👩‍💼 Solicitor created: %s (Profession: %s)",
                solicitor,
                profession)

            # Prepare admin email with solicitor/conveyancer license details
            admin_email_body = (
                f"A new {profession} has registered:\n\n"
                f"Name: {first_name} {last_name}\n"
                f"Firm Name: {firm_name}\n"
                f"Email: {email}\n"
                f"Office Phone: {office_phone}\n"
                f"Mobile Phone: {mobile_phone}\n"
                f"Profession: {profession}\n"
            )
            if profession == "solicitor":
                admin_email_body += f"Law Society Number: {law_society_number}\n"
            elif profession == "conveyancer":
                admin_email_body += f"Conveyancer License: {conveyancer_license_number}\n"

            # Send email to admin
            send_mail(
                subject="New Solicitor Registration Submitted",
                message=admin_email_body,
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=['info@onestoplegal.com.au'],
                fail_silently=False,
            )
            logger.info("✅ Email sent to admin successfully")

            # Send confirmation email to user
            send_mail(
                subject="Registration Submitted - SettleX",
                message=f"Dear {first_name} {last_name},\n\n"
                        f"Thank you for registering with SettleX. Your registration is currently pending approval.\n"
                        f"You will receive an email once your account has been activated.\n\n"
                        f"Regards,\nSettleX Team",
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[email],
                fail_silently=False,
            )
            logger.info("✅ Confirmation email sent to user successfully")

            messages.success(
                request, "Registration submitted! Please log in to continue.")
            logger.info("✅ Registration successful for: %s", user.username)
            return redirect('settlements_app:login')  # Redirect to login page

        except Exception as e:
            messages.error(request, f"Error during registration: {str(e)}")
            logger.error("🚨 Registration error: %s", str(e))
            return render(request,
                          'settlements_app/register.html',
                          {'message': 'Register an Account'})

    return render(request, 'settlements_app/register.html',
                  {'existing_firm': firm, 'page_title': 'Register an Account'})



def new_instruction(request):
    try:
        if request.method == 'POST':
            logger.info("✅ Received POST request for new instruction.")

            file_reference = request.POST.get('file_reference', '').strip()
            if not file_reference:
                messages.error(request, "File Reference is required.")
                return redirect('settlements_app:new_instruction')

            # Check for duplicate file_reference
            if Instruction.objects.filter(file_reference=file_reference).exists():
                messages.error(request, "This file reference already exists. Please use a different one.")
                return redirect('settlements_app:new_instruction')

            transaction_type = request.POST.get('instruction_type', '').strip()
            title_reference = request.POST.get('title_reference', '').strip()

            if transaction_type in ['purchase', 'sale']:
                date_value = request.POST.get('settlement_date', '').strip()
            else:
                date_value = request.POST.get('lodgement_date', '').strip()

            if not date_value:
                messages.error(request, "A settlement or lodgement date is required.")
                return redirect('settlements_app:new_instruction')

            try:
                settlement_date = datetime.strptime(date_value, '%Y-%m-%d').date()
            except ValueError:
                messages.error(request, "Invalid date format. Use YYYY-MM-DD.")
                return redirect('settlements_app:new_instruction')

            # Transaction address fields
            street_number = request.POST.get('transaction_street_number', '').strip()
            street_name = request.POST.get('transaction_street_name', '').strip()
            suburb = request.POST.get('transaction_suburb', '').strip()
            state = request.POST.get('transaction_state', '').strip()
            postcode = request.POST.get('transaction_postcode', '').strip()
            property_type = request.POST.get('property_type', '').strip()

            if not (street_number and street_name and suburb and state and postcode):
                messages.error(request, "Complete transaction address is required.")
                return redirect('settlements_app:new_instruction')

            if not title_reference:
                messages.error(request, "Title reference(s) are required.")
                return redirect('settlements_app:new_instruction')

            solicitor = getattr(request.user, 'solicitor', None)
            if not solicitor:
                messages.error(request, "You must be a registered solicitor to submit instructions.")
                return redirect('home')

            # Optional financial fields
            purchase_price_raw = request.POST.get('purchase_price', '').strip()
            deposit_raw = request.POST.get('deposit', '').strip()

            purchase_price = None
            if purchase_price_raw:
                try:
                    purchase_price = Decimal(purchase_price_raw.replace(',', ''))
                except InvalidOperation:
                    messages.error(request, "Invalid purchase price format.")
                    return redirect('settlements_app:new_instruction')

            # Allow deposit to be any string (amount or percentage)
            deposit = deposit_raw if deposit_raw else None

            with transaction.atomic():
                instruction = Instruction.objects.create(
                    solicitor=solicitor,
                    file_reference=file_reference,
                    settlement_type=transaction_type,
                    settlement_date=settlement_date,
                    title_reference=title_reference,
                    transaction_street_number=street_number,
                    transaction_street_name=street_name,
                    transaction_suburb=suburb,
                    transaction_state=state,
                    transaction_postcode=postcode,
                    property_type=property_type,
                    purchase_price=purchase_price,
                    deposit=deposit
                )

                # Client section
                client_type = request.POST.get('client')
                if client_type == 'individual':
                    individuals = []
                    num_individuals = int(request.POST.get('num_individuals', 0))
                    for i in range(1, num_individuals + 1):
                        name = request.POST.get(f'individual_name_{i}', '').strip()
                        dob = request.POST.get(f'individual_dob_{i}', '').strip()
                        email = request.POST.get(f'individual_email_{i}', '').strip()
                        mobile = request.POST.get(f'individual_mobile_{i}', '').strip()
                        address_line = request.POST.get(f'individual_address_{i}', '').strip()
                        suburb_line = request.POST.get(f'individual_suburb_{i}', '').strip()
                        state_line = request.POST.get(f'individual_state_{i}', '').strip()
                        postcode_line = request.POST.get(f'individual_postcode_{i}', '').strip()
                        full_address = f"{address_line}, {suburb_line} {state_line} {postcode_line}"
                        individuals.append({
                            'name': name,
                            'email': email,
                            'mobile': mobile,
                            'address': full_address
                        })

                    instruction.purchaser_name = "; ".join([i['name'] for i in individuals if i['name']])
                    instruction.purchaser_mobile = "; ".join([i['mobile'] for i in individuals if i['mobile']])
                    instruction.purchaser_email = "; ".join([i['email'] for i in individuals if i['email']])
                    instruction.purchaser_address = "; ".join([i['address'] for i in individuals if i['address']])

                elif client_type == 'company':
                    instruction.purchaser_name = request.POST.get('company_name', '').strip()
                    instruction.purchaser_mobile = request.POST.get('company_abn', '').strip()
                    instruction.purchaser_email = request.POST.get('company_acn', '').strip()

                    company_street = request.POST.get('company_street', '').strip()
                    company_suburb = request.POST.get('company_suburb', '').strip()
                    company_state = request.POST.get('company_state', '').strip()
                    company_postcode = request.POST.get('company_postcode', '').strip()
                    full_company_address = f"{company_street}, {company_suburb} {company_state} {company_postcode}"
                    instruction.purchaser_address = full_company_address

                instruction.save()

                # Send summary email
                email_subject = f"New Instruction {instruction.file_reference}"
                email_body = (
                    f"File Reference: {instruction.file_reference}\n"
                    f"Solicitor: {solicitor.instructing_solicitor}\n"
                    f"Firm: {solicitor.firm.name if solicitor.firm else ''}\n"
                    f"Settlement Type: {instruction.settlement_type}\n"
                    f"Title Reference: {instruction.title_reference}\n"
                    f"Settlement Date: {instruction.settlement_date}"
                )
                try:
                    send_mail(
                        subject=email_subject,
                        message=email_body,
                        from_email=settings.DEFAULT_FROM_EMAIL,
                        recipient_list=[solicitor.user.email],
                        fail_silently=False,
                    )
                    messages.success(request, "Instruction created successfully!")
                except Exception as email_error:
                    logger.error("Email send failed: %s", email_error)
                    messages.success(request, "Instruction created successfully, but failed to send email notification.")

            logger.info(f"✅ Instruction created successfully: {instruction.file_reference}")
            return redirect('settlements_app:my_transactions')

        return render(request, 'settlements_app/new_instruction.html', {'page_title': 'Create New Instruction'})

    except Exception as e:
        logger.error(traceback.format_exc())
        messages.error(request, f"An unexpected error occurred: {str(e)}")
        return redirect('settlements_app:new_instruction')


@otp_required
def my_transactions(request):
    try:
        solicitor = getattr(request.user, 'solicitor', None)
        logger.debug("👤 Solicitor for user %s: %s", request.user.username, solicitor)

        if not solicitor:
            messages.warning(
                request,
                "Your account is not fully registered. Please complete your solicitor profile.")
            logger.warning("⚠️ No solicitor profile found for user: %s", request.user.username)
            transactions = []
        elif not solicitor.firm:
            messages.warning(
                request,
                "Your solicitor profile is not linked to a firm. Please update your profile.")
            logger.warning("⚠️ No firm linked to solicitor: %s", solicitor)
            transactions = []
        else:
            # ✅ Fetch instructions related to the solicitor's firm
            transactions = Instruction.objects.filter(
                solicitor__firm=solicitor.firm
            ).order_by('-settlement_date')
            logger.info("📄 Retrieved %d transactions for firm: %s", transactions.count(), solicitor.firm)

        chat_messages = ChatMessage.objects.filter(
            recipient=request.user
        ).order_by("timestamp")
        logger.debug("💬 Loaded %d chat messages for user: %s", chat_messages.count(), request.user.username)

        # Set enable_chat to True for this page
        enable_chat = True

    except Exception as e:
        logger.error("🚨 Error loading transactions: %s", str(e))
        messages.error(request, "An error occurred while loading your transactions.")
        transactions = []
        chat_messages = []
        enable_chat = False  # If there's an error, chat is disabled

    return render(request, 'settlements_app/my_transactions.html', {
        'transactions': transactions,
        'chat_messages': chat_messages,
        'enable_chat': enable_chat,  # Pass enable_chat to the template
    })

def upload_documents(request):
    """Allows solicitors to upload documents for any instruction within their firm."""
    solicitor = getattr(request.user, 'solicitor', None)

    if not solicitor or not solicitor.firm:
        messages.error(
            request,
            "You must be a registered solicitor with a firm to access this page.")
        return redirect('home')  # Assuming 'home' is global and not namespaced

    try:
        instructions = Instruction.objects.filter(
            solicitor__firm=solicitor.firm).order_by('-settlement_date')

        preselected_instruction = None
        settlement_id = request.GET.get('settlement_id') or request.POST.get('instruction_id')

        if settlement_id:
            try:
                preselected_instruction = get_object_or_404(
                    Instruction,
                    id=settlement_id,
                    solicitor__firm=solicitor.firm
                )
            except Exception as e:
                logger.error(f"❌ Error finding settlement instruction: {e}")
                messages.error(request, "Invalid settlement ID.")
                # ✅ Namespaced
                return redirect('settlements_app:upload_documents')

        if request.method == 'POST' and request.FILES.get('document'):
            if not preselected_instruction:
                messages.error(
                    request, "No valid instruction selected for document upload.")
                # ✅ Namespaced
                return redirect('settlements_app:upload_documents')

            uploaded_file = request.FILES['document']
            document_type = request.POST.get(
                'document_type', 'contract').strip()
            document_name = request.POST.get(
                'document_name', uploaded_file.name).strip()

            DocumentModel.objects.create(
                instruction=preselected_instruction,
                name=document_name,
                file=uploaded_file,
                document_type=document_type
            )

            messages.success(
                request, f"File '{document_name}' uploaded successfully!")
            return redirect(
                'settlements_app:view_settlement',
                settlement_id=preselected_instruction.id)  # ✅ Namespaced

    except Exception as e:
        logger.error(f"❌ Error uploading documents: {traceback.format_exc()}")
        messages.error(
            request,
            "An unexpected error occurred while uploading documents.")

    return render(request, 'settlements_app/upload_documents.html', {
        'instructions': instructions,
        'preselected_instruction': preselected_instruction,
        'page_title': 'Upload Documents'
    })

# ✅ Admin Solicitor Dashboard


def solicitor_dashboard(request):
    """Displays all firms and their settlements for admin monitoring."""
    if not request.user.is_superuser:
        messages.error(request, "Access denied.")
        return redirect('home')

    try:
        firms = Solicitor.objects.all()
    except Exception as e:
        logger.error(f"❌ Error fetching solicitor firms: {e}")
        messages.error(
            request,
            "An error occurred while retrieving solicitor data.")
        firms = []

    return render(request, 'settlements_app/solicitor_dashboard.html',
                  {'firms': firms, 'page_title': 'Solicitor Dashboard'})

# ✅ Edit Instruction View


def edit_instruction(request, id):
    """Edit an existing instruction."""
    solicitor = getattr(request.user, 'solicitor', None)
    if not solicitor:
        messages.error(
            request,
            "You must be a registered solicitor to edit an instruction.")
        return redirect('home')

    instruction = get_object_or_404(Instruction, id=id, solicitor=solicitor)

    if request.method == "POST":
        form = InstructionForm(request.POST, instance=instruction)
        if form.is_valid():
            form.save()
            messages.success(request, "Instruction updated successfully!")
            return redirect('settlements_app:my_transactions')
    else:
        form = InstructionForm(instance=instruction)

    return render(request, 'settlements_app/edit_instruction.html', {
        'form': form,
        'instruction': instruction,
        'page_title': 'Edit Instruction'
    })

# ✅ Delete Instruction View


def delete_instruction(request, id):
    """Deletes an instruction"""
    solicitor = getattr(request.user, 'solicitor', None)
    if not solicitor:
        messages.error(
            request,
            "You must be a registered solicitor to delete an instruction.")
        return redirect('home')

    instruction = get_object_or_404(Instruction, id=id, solicitor=solicitor)

    if request.method == 'POST':
        instruction.delete()
        messages.success(request, "Instruction deleted successfully!")
        return redirect('settlements_app:my_transactions')

    return render(request, 'settlements_app/delete_instruction.html', {
        'instruction': instruction,
        'page_title': 'Delete Instruction'
    })



from decimal import Decimal, InvalidOperation
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, redirect, render
from .models import Instruction, PaymentDirection
from .forms import PaymentDirectionForm, PaymentDirectionLineItemForm


@login_required
def payment_direction(request, instruction_id):
    """View or add/edit payment direction and its line items by type."""
    solicitor = getattr(request.user, 'solicitor', None)
    if not solicitor or not solicitor.firm:
        messages.error(
            request,
            "You must be a registered solicitor with a firm to access this page.")
        return redirect('settlements_app:home')

    instruction = get_object_or_404(
        Instruction,
        id=instruction_id,
        solicitor__firm=solicitor.firm,
    )

    payment = PaymentDirection.objects.filter(instruction=instruction).first()
    if not payment:
        payment = PaymentDirection.objects.create(instruction=instruction)

    # Separate line items by direction type
    purchaser_line_items = payment.line_items.filter(direction_type='purchaser')
    vendor_line_items = payment.line_items.filter(direction_type='vendor')

    # Totals
    total_purchaser_amount = sum(item.amount for item in purchaser_line_items)
    total_vendor_amount = sum(item.amount for item in vendor_line_items)

    # Parse deposit input
    purchase_price = instruction.purchase_price or Decimal('0')
    deposit_raw = instruction.deposit or ''
    adjustments = instruction.adjustments or Decimal('0')
    deposit_amount = Decimal('0')

    try:
        clean = deposit_raw.strip().replace(',', '').replace('$', '')
        if clean.endswith('%'):
            percent = Decimal(clean.rstrip('%').strip())
            deposit_amount = (purchase_price * percent) / Decimal('100')
        else:
            deposit_amount = Decimal(clean)
    except (InvalidOperation, ValueError):
        deposit_amount = Decimal('0')  # Fallback for invalid or text-based deposits

    # Financial calculations
    balance_owing_to_vendor = purchase_price - deposit_amount - adjustments
    total_amount_purchaser_has_to_pay = balance_owing_to_vendor + total_purchaser_amount
    funds_available_to_settle = payment.funds_available_to_settle or Decimal('0')

    # Check if we're ready to settle
    settlement_ready = funds_available_to_settle == total_amount_purchaser_has_to_pay

    # Determine shortfall or surplus
    difference = Decimal('0')
    surplus = Decimal('0')
    if funds_available_to_settle < total_amount_purchaser_has_to_pay:
        difference = total_amount_purchaser_has_to_pay - funds_available_to_settle
    elif funds_available_to_settle > total_amount_purchaser_has_to_pay:
        surplus = funds_available_to_settle - total_amount_purchaser_has_to_pay

    direction_tables = [
        {
            'label': "Purchaser Payment Directions",
            'items': purchaser_line_items,
            'total': total_purchaser_amount
        },
        {
            'label': "Vendor Payment Directions",
            'items': vendor_line_items,
            'total': total_vendor_amount
        }
    ]

    form = PaymentDirectionForm(request.POST or None, instance=payment)
    line_item_form = PaymentDirectionLineItemForm(request.POST or None)

    if request.method == 'POST':
        if 'save_line_item' in request.POST and line_item_form.is_valid():
            item = line_item_form.save(commit=False)
            item.payment_direction = payment
            item.save()
            messages.success(request, 'Payment direction line item added successfully.')
            return redirect('settlements_app:payment_direction', instruction_id=instruction.id)

        elif 'save_main' in request.POST or 'save_all' not in request.POST:
            if form.is_valid():
                form.save()
                messages.success(request, 'Payment direction saved successfully.')
                return redirect('settlements_app:payment_direction', instruction_id=instruction.id)

        elif 'save_all' in request.POST:
            for item in payment.line_items.all():
                item.category = request.POST.get(f'category_{item.id}', item.category)
                item.bank_name = request.POST.get(f'bank_name_{item.id}', item.bank_name)
                item.account_name = request.POST.get(f'account_name_{item.id}', item.account_name)
                item.account_details = request.POST.get(f'account_details_{item.id}', item.account_details)
                item.amount = request.POST.get(f'amount_{item.id}', item.amount)
                item.save()
            messages.success(request, 'All payment direction line items saved successfully.')
            return redirect('settlements_app:payment_direction', instruction_id=instruction.id)

    return render(
        request,
        'settlements_app/payment_direction_form.html',
        {
            'form': form,
            'instruction': instruction,
            'line_item_form': line_item_form,
            'direction_tables': direction_tables,
            'purchase_price': purchase_price,
            'deposit': deposit_raw,
            'deposit_amount': deposit_amount,
            'adjustments': adjustments,
            'balance_owing_to_vendor': balance_owing_to_vendor,
            'total_amount_purchaser_has_to_pay': total_amount_purchaser_has_to_pay,
            'funds_available_to_settle': funds_available_to_settle,
            'settlement_ready': settlement_ready,
            'difference': difference,
            'surplus': surplus,
        }
    )



@login_required
@csrf_protect  # CSRF protection for the delete action
def delete_line_item(request, item_id):
    """Delete a payment direction line item."""
    if request.method == 'POST':
        # Fetch item or return 404 if not found
        item = get_object_or_404(PaymentDirectionLineItem, id=item_id)

        try:
            item.delete()
            return JsonResponse({'status': 'success'})
        except Exception as e:
            # Catch errors if any
            return JsonResponse({'status': 'error', 'message': f'Error deleting item: {str(e)}'})

    return JsonResponse({'status': 'error', 'message': 'Invalid request method'})



@login_required
@otp_required
@require_POST
def edit_line_item(request):
    """Edit a specific line item via AJAX."""
    item_id = request.POST.get('item_id')
    item = get_object_or_404(PaymentDirectionLineItem, id=item_id)

    solicitor = getattr(request.user, 'solicitor', None)
    if not solicitor or not solicitor.firm or item.payment_direction.instruction.solicitor.firm != solicitor.firm:
        return JsonResponse({'status': 'error', 'message': 'Unauthorized'}, status=403)

    form = PaymentDirectionLineItemForm(request.POST, instance=item)
    if form.is_valid():
        form.save()
        return JsonResponse({'status': 'success', 'message': 'Line item updated'})
    else:
        logger.error("Edit Line Item Form Errors: %s", form.errors)
        return JsonResponse({'status': 'error', 'errors': form.errors}, status=400)




@login_required
def list_payment_directions(request):
    """List all matters with links to their Payment Direction forms."""
    solicitor = getattr(request.user, 'solicitor', None)
    if not solicitor or not solicitor.firm:
        messages.error(request, "You must be a registered solicitor with a firm to access this page.")
        return redirect('settlements_app:home')

    instructions = Instruction.objects.filter(solicitor__firm=solicitor.firm)

    return render(request, 'settlements_app/list_payment_directions.html', {
        'instructions': instructions,
    })
# ✅ View Settlement Details


def view_transaction(request, transaction_id):
    """View transaction details and related documents for the user's firm."""
    solicitor = getattr(request.user, 'solicitor', None)

    if not solicitor or not solicitor.firm:
        messages.error(
            request,
            "You must be a registered solicitor with a firm to access this page.")
        return redirect('settlements_app:home')

    try:
        # ✅ Ensure solicitors from the same firm can see each other's transactions
        transaction = get_object_or_404(
            Instruction,
            id=transaction_id,
            solicitor__firm=solicitor.firm
        )

        # ✅ Fetch all documents linked to this transaction
        documents = DocumentModel.objects.filter(instruction=transaction)

    except Exception as e:
        logger.error(f"❌ Error loading transaction details: {e}")
        messages.error(
            request,
            "An error occurred while retrieving transaction details.")
        return redirect('settlements_app:my_transactions')

    context = {
        'transaction': transaction,
        'documents': documents,
        'preselected_instruction': transaction,
        'page_title': 'View Transaction',
    }

    return render(request, 'settlements_app/view_transactions.html', context)



# ✅ Set up logger
logger = logging.getLogger(__name__)

# Set Brisbane timezone
BRISBANE_TZ = pytz.timezone("Australia/Brisbane")

# Get current Brisbane time


def get_brisbane_time():
    return now().astimezone(BRISBANE_TZ)


def long_poll_messages(request):
    """Fetch full chat history (both sent & received messages) for the logged-in user."""
    user = request.user
    logger.info(
        f"📩 Long poll request from user: {user} (ID: {user.id}) at {now()} - Poll cycle start")

    try:
        logger.debug("Fetching messages...")
        messages = ChatMessage.objects.filter(
            Q(sender=user) | Q(recipient=user))
        messages = messages.filter(
            timestamp__gte=now() -
            timedelta(
                days=7)).order_by("timestamp")

        total_messages = messages.count()
        logger.info(f"📬 Total messages fetched for {user}: {total_messages}")

        if total_messages == 0:
            logger.info("No messages found, returning empty response.")
            return JsonResponse(
                {"messages": [], "status": "success"}, status=200)

        messages_data = []
        for msg in messages:
            try:
                timestamp_str = localtime(
                    msg.timestamp, BRISBANE_TZ).strftime("%d %b %Y, %I:%M %p")
                is_read_value = msg.is_read
                role = "sender" if msg.sender == user else "recipient" if msg.recipient == user else "other"
                logger.debug(
                    f"Poll at {now()}: Message {msg.id} - is_read={is_read_value}, timestamp={timestamp_str}, sender={msg.sender.username}, recipient={msg.recipient.username}, user_role={role}")
                message_dict = {
                    "id": msg.id,
                    "sender_name": msg.sender.get_full_name() or msg.sender.username,
                    "sender_username": msg.sender.username,  # New field for direct comparison
                    "recipient_name": msg.recipient.get_full_name() or msg.recipient.username,
                    "message": msg.message,
                    "timestamp": timestamp_str,
                    "is_read": is_read_value,
                    "user_role": role,
                    "file_url": msg.file.url if msg.file else None  # Include file URL if exists
                }
                messages_data.append(message_dict)
            except Exception as e:
                logger.error(
                    f"Error processing message {msg.id}: {str(e)}",
                    exc_info=True)
                continue

        logger.info(
            f"Poll cycle completed at {now()} - Returning successful response with {len(messages_data)} messages.")
        return JsonResponse(
            {"messages": messages_data, "status": "success"}, status=200)

    except Exception as e:
        error_details = traceback.format_exc()
        logger.error(
            f"❌ ERROR in long_poll_messages for {user}: {e}\n{error_details}",
            exc_info=True)
        return JsonResponse({"status": "error",
                             "message": f"Could not fetch messages: {str(e)}"},
                            status=500)


def send_message(request):
    """Send a chat message to another user."""
    logger.info(
        f"User authenticated: {request.user.is_authenticated}, User: {request.user}")

    if not request.user.is_authenticated:
        logger.warning("Unauthenticated user attempted to send a message")
        return JsonResponse(
            {"status": "error", "message": "Authentication required"},
            status=403,
        )

    if request.method == "POST":
        try:
            logger.info(
                f"Received POST request from user {request.user.username}: POST={request.POST}, FILES={request.FILES}")
            message_text = request.POST.get("message", "").strip()
            file = request.FILES.get("file")
            recipient_id = request.POST.get("recipient")
            logger.debug(
                f"Processing message: text='{message_text}', file={file}, recipient_id={recipient_id}")

            if not message_text and not file:
                logger.warning("No message text or file provided")
                return JsonResponse(
                    {"status": "error", "message": "Message text or file required"}, status=400)

            if not recipient_id:
                logger.error("No recipient ID provided in POST data")
                return JsonResponse(
                    {"status": "error", "message": "Recipient required"}, status=400)

            try:
                recipient = User.objects.get(id=recipient_id)
            except User.DoesNotExist:
                logger.error(f"Recipient with ID {recipient_id} not found")
                return JsonResponse(
                    {"status": "error", "message": "Invalid recipient"}, status=404)

            message = ChatMessage(
                sender=request.user,
                recipient=recipient,
                message=message_text if message_text else None,
                file=file,
                is_read=False
            )
            message.save()
            logger.info(
                f"Message saved successfully: ID={message.id}, initial_is_read={message.is_read}")

            response_data = {
                "status": "success",
                "message": "Message sent",
                "id": message.id,
                "is_read": message.is_read,
                "sender_name": request.user.get_full_name() or request.user.username,
                "sender_username": request.user.username,
                "recipient_name": recipient.get_full_name() or recipient.username,
                "timestamp": localtime(
                    message.timestamp,
                    BRISBANE_TZ).strftime("%d %b %Y, %I:%M %p"),
                "file_url": message.file.url if message.file else None}
            logger.debug(f"Message response: {response_data}")
            return JsonResponse(response_data)
        except Exception as e:
            logger.error(f"Error in send_message: {str(e)}", exc_info=True)
            return JsonResponse(
                {"status": "error", "message": str(e)}, status=500)
    else:
        logger.warning(f"Invalid request method: {request.method}")
        return JsonResponse(
            {"status": "error", "message": "Invalid request method"}, status=400)



def check_new_messages(request):
    """Check for new messages for the logged-in user and optionally mark them as read."""
    user = request.user
    logger.info(
        f"🔍 Checking new messages for user: {user} (ID: {user.id}) at {now()}")

    try:
        # Fetch unread messages where the user is the recipient
        unread_messages = ChatMessage.objects.filter(
            recipient=user, is_read=False).exclude(
            sender=user)
        total_unread = unread_messages.count()
        logger.info(f"📬 Found {total_unread} unread messages for {user}")

        if total_unread == 0:
            logger.info("No new messages found.")
            return JsonResponse(
                {"status": "success", "new_messages": 0}, status=200)

        # Optionally mark messages as read (if intended by original design)
        updated_count = unread_messages.update(is_read=True)
        logger.info(f"✅ Marked {updated_count} messages as read for {user}")

        return JsonResponse(
            {"status": "success", "new_messages": updated_count}, status=200)

    except Exception as e:
        logger.error(
            f"❌ Error checking new messages for {user}: {str(e)}",
            exc_info=True)
        return JsonResponse({"status": "error", "message": str(e)}, status=500)

# ✅ Reply to a Message


def reply_view(request, message_id):
    """Handle message replies from admin."""
    message = get_object_or_404(ChatMessage, id=message_id)

    if request.method == "POST":
        reply_text = request.POST.get("reply_message", "").strip()

        if reply_text:
            # ✅ Ensure recipient is the original sender (if valid)
            if message.sender != request.user:
                recipient = message.sender  # Correct recipient
            else:
                # Avoid replying to Settlex (Admin)
                recipient = message.recipient

            # ✅ Save the reply correctly
            ChatMessage.objects.create(
                sender=request.user,  # Always the admin
                recipient=recipient,  # Ensure this is the actual user
                message=reply_text,
                timestamp=now(),
                is_read=False
            )

            messages.success(request, "Reply sent successfully!")
            return HttpResponseRedirect(
                reverse("admin:settlements_app_chatmessage_changelist"))

    return render(request, "admin/chat_reply.html", {
        "message": message,
        "subtitle": "Replying to Chat Message"
    })


@csrf_protect  # Protects this view with CSRF validation
def mark_messages_read(request):
    """Mark messages as read when a user views them."""
    if request.method != "POST":
        return JsonResponse(
            {"status": "error", "message": "Invalid request method"}, status=400)

    try:
        data = json.loads(request.body)  # Read JSON data properly
        message_ids = data.get("message_ids", [])

        if not message_ids:
            return JsonResponse(
                {"status": "error", "message": "No message IDs provided"}, status=400)

        # Convert IDs to integers (skip invalid ones)
        message_ids = [int(msg_id)
                       for msg_id in message_ids if str(msg_id).isdigit()]

        # Update messages if user is authenticated
        messages = ChatMessage.objects.filter(
            id__in=message_ids, recipient=request.user, is_read=False)
        messages.update(is_read=True)  # Efficient bulk update

        return JsonResponse(
            {"status": "success", "updated": len(messages)}, status=200)

    except json.JSONDecodeError:
        return JsonResponse(
            {"status": "error", "message": "Invalid JSON data"}, status=400)

    except Exception as e:
        return JsonResponse({"status": "error",
                             "message": f"Could not mark messages read: {str(e)}"},
                            status=500)


def check_typing_status(request):
    is_typing = request.GET.get('admin_typing', 'false').lower() == 'true'
    return JsonResponse({"is_typing": is_typing})


def upload_chat_file(request):
    if request.method == "POST" and request.FILES.get("file"):
        file = request.FILES["file"]
        from django.core.files.storage import default_storage
        file_name = default_storage.save(f"chat_files/{file.name}", file)
        file_url = default_storage.url(file_name)
        return JsonResponse({"status": "success", "file_url": file_url})
    return JsonResponse(
        {"status": "error", "message": "No file uploaded"}, status=400)


def delete_message(request):
    if request.method == "POST":
        data = json.loads(request.body)
        message_id = data.get("message_id")
        try:
            message = ChatMessage.objects.get(
                id=message_id, sender=request.user)
            message.delete()
            return JsonResponse({"status": "success"})
        except ChatMessage.DoesNotExist:
            return JsonResponse(
                {"status": "error", "message": "Message not found or unauthorized"}, status=403)
    return JsonResponse(
        {"status": "error", "message": "Invalid request"}, status=400)


def calculate_adjustment(full_amount, period_start, period_end, adj_date):
    """Pro-rata adjustment: seller reimbursed for unused days after settlement."""
    if not (period_start and period_end and adj_date):
        return Decimal('0.00'), 0, Decimal('0.00')

    try:
        total_days = (period_end - period_start).days + 1
        buyer_days = (period_end - adj_date).days
        buyer_days = max(0, min(buyer_days, total_days))

        daily_rate = full_amount / total_days if total_days > 0 else Decimal('0.00')
        adjustment_amount = (daily_rate * buyer_days).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        return daily_rate.quantize(Decimal('0.01')), buyer_days, adjustment_amount
    except Exception:
        logger.exception("Failed to calculate adjustment")
        return Decimal('0.00'), 0, Decimal('0.00')

from decimal import Decimal, ROUND_HALF_UP
from django.contrib import messages
from django.shortcuts import render, get_object_or_404, redirect
from datetime import timedelta
from .models import Instruction
from .forms import RatesAdjustmentForm
import logging

logger = logging.getLogger(__name__)

def rates_adjustment_view(request, instruction_id):
    instruction = get_object_or_404(Instruction, id=instruction_id)

    daily_rate = None
    buyer_days = None
    calculated_amount = None

    if request.method == 'POST':
        form = RatesAdjustmentForm(request.POST)
        if form.is_valid():
            adj = form.save(commit=False)

            # Ensure required fields are provided
            if adj.period_from and adj.period_to and adj.total_amount and instruction.settlement_date:
                try:
                    total_days = (adj.period_to - adj.period_from).days + 1
                    buyer_start = instruction.settlement_date + timedelta(days=1)
                    buyer_days = (adj.period_to - buyer_start).days + 1

                    # Bounds check
                    buyer_days = max(0, min(buyer_days, total_days))

                    daily_rate = (adj.total_amount / Decimal(total_days)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
                    calculated_amount = (daily_rate * Decimal(buyer_days)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

                    adj.amount = calculated_amount
                except Exception as e:
                    logger.exception("Error calculating rates adjustment: %s", e)
                    messages.error(request, "An error occurred while calculating the adjustment.")
                    adj.amount = Decimal('0.00')
            else:
                adj.amount = Decimal('0.00')

            adj.instruction = instruction
            adj.save()

            messages.success(request, "Rates adjustment saved successfully.")
            return redirect('settlements_app:rates_adjustment', instruction_id=instruction.id)
    else:
        form = RatesAdjustmentForm()

    return render(request, 'settlements_app/rates_adjustment.html', {
        'form': form,
        'instruction': instruction,
        'daily_rate': daily_rate,
        'buyer_days': buyer_days,
        'calculated_amount': calculated_amount,
    })



def settlement_calculator(request):
    result = None

    body_corp_levies = [
        {'label': 'Admin Fund', 'name': 'admin'},
        {'label': 'Sinking Fund', 'name': 'sinking'},
        {'label': 'Insurance', 'name': 'insurance'},
        {'label': 'Special Levy', 'name': 'special'},
    ]


    calc_form = SettlementCalculatorForm(request.POST or None)

    if request.method == 'POST' and calc_form.is_valid():
        try:
            logger.info("Starting settlement calculation.")

            settlement_date = calc_form.cleaned_data['settlement_date']
            adjustment_date = calc_form.cleaned_data['adjustment_date']
            contract_date = calc_form.cleaned_data.get('contract_date')
            settlement_place = calc_form.cleaned_data.get('settlement_place')
            settlement_time = calc_form.cleaned_data.get('settlement_time')
            contract_price = calc_form.cleaned_data.get('contract_price') or 0
            deposit = calc_form.cleaned_data.get('deposit') or 0
            release_mortgage_fee = calc_form.cleaned_data.get('release_mortgage_fee') or 0
            registration_fee = calc_form.cleaned_data.get('registration_fee') or 0
            pexa_fee = calc_form.cleaned_data.get('pexa_fee') or 0

            # Calculate adjustments based on the form data
            council_adj = calculate_adjustment(
                calc_form.cleaned_data.get('council_amount', 0),
                calc_form.cleaned_data.get('council_start'),
                calc_form.cleaned_data.get('council_end'),
                adjustment_date
            )

            water_adj = calculate_adjustment(
                calc_form.cleaned_data.get('water_amount', 0),
                calc_form.cleaned_data.get('water_start'),
                calc_form.cleaned_data.get('water_end'),
                adjustment_date
            )


            bodycorp_adjustments = {}
            for levy in body_corp_levies:
                bodycorp_adjustments[f"bodycorp_{levy['name']}"] = calculate_adjustment(
                    calc_form.cleaned_data.get(f"{levy['name']}_amount", 0),
                    calc_form.cleaned_data.get(f"{levy['name']}_start"),
                    calc_form.cleaned_data.get(f"{levy['name']}_end"),
                    adjustment_date
                )

            total_adjustments = council_adj + water_adj + sum(bodycorp_adjustments.values())
            balance_at_settlement = (
                (contract_price or 0)
                - deposit
                + release_mortgage_fee
                + total_adjustments
                + registration_fee
                + pexa_fee
            )

            result = {
                'settlement_date': settlement_date,
                'adjustment_date': adjustment_date,
                'contract_date': contract_date,
                'settlement_place': settlement_place,
                'settlement_time': settlement_time,
                'contract_price': contract_price,
                'deposit': deposit,
                'release_mortgage_fee': release_mortgage_fee,
                'registration_fee': registration_fee,
                'pexa_fee': pexa_fee,
                'council_adj': council_adj,
                'water_adj': water_adj,
                **bodycorp_adjustments,
                'total_adjustments': total_adjustments,
                'balance_at_settlement': balance_at_settlement
            }

            request.session['settlement_data'] = {
                **{key: str(value) for key, value in result.items()},
                'body_corp_levies': body_corp_levies,
            }

            logger.info("Settlement calculation completed and data stored in session.")

            return redirect('settlements_app:settlement_statement')

        except Exception as e:
            logger.error(f"Error occurred in settlement calculation: {str(e)}")
            messages.error(request, f"Error: {str(e)}")

    return render(request, 'settlements_app/settlement_calculator.html', {
        'adjustments_form': calc_form,
        'result': result,
        'body_corp_levies': body_corp_levies
    })

def settlement_statement(request):
    # Retrieve data from the session
    data = request.session.get('settlement_data')

    # Log the session data for debugging
    if not data:
        logger.error("No settlement data found in session.")
        return redirect('settlements_app:settlement_calculator')  # Redirect back to the calculator if no data is found

    logger.info(f"Settlement Data found in session: {data}")

    # Convert values back to float for display
    for key in [
        'contract_price',
        'deposit',
        'release_mortgage_fee',
        'registration_fee',
        'pexa_fee',
        'council_adj',
        'water_adj',
        'bodycorp_admin',
        'bodycorp_sinking',
        'bodycorp_insurance',
        'bodycorp_special',
        'total_adjustments',
        'balance_at_settlement',
    ]:
        data[key] = float(data.get(key, 0))

    # You can now render the statement template
    return render(request, 'settlements_app/settlement_statement.html', {'data': data})

def settlement_statement_word(request):
    data = request.session.get('settlement_data')
    if not data:
        return redirect('settlements_app:settlement_calculator')

    # Cast data values to float if necessary
    for key in [
        'contract_price',
        'deposit',
        'release_mortgage_fee',
        'registration_fee',
        'pexa_fee',
        'council_adj',
        'water_adj',
        'bodycorp_admin',
        'bodycorp_sinking',
        'bodycorp_insurance',
        'bodycorp_special',
        'total_adjustments',
        'balance_at_settlement',
    ]:
        data[key] = float(data.get(key, 0))
    document = Document()
    document.add_heading('Settlement Statement', level=1)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Value'

    for key, value in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key.replace('_', ' ').title()
        row_cells[1].text = str(value)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return FileResponse(buffer, as_attachment=True, filename="settlement_statement.docx")
